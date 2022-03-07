import datetime
import glob
import json
import logging.config
import mammoth
import numpy as np
import os
import secrets
import shutil
import urllib3
from PIL import Image
from docx import Document
from docx.shared import Inches
from firewall import Firewall
from flask import Flask, flash, request, redirect, render_template, url_for, send_file, make_response, Response, session
from flask_bcrypt import Bcrypt
from flask_session import Session
from functools import wraps
from helpers import PanHelpers
from htmldocx import HtmlToDocx
from paloalto import PaloFirewall, PaloPanorama
from panorama import Panorama
from swift_doc import SwiftDoc
from tempfile import mkdtemp
from tsf_handler import PanTSFHandler
from werkzeug.utils import secure_filename

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

logging.config.fileConfig(
    "logging.conf",
    defaults={"logfilename": "SwiftDoc.log"},
    disable_existing_loggers=False,
)
logger = logging.getLogger("__name__")
# context = ssl.SSLContext()
# context.load_cert_chain('laragon.crt', 'laragon.key')
app = Flask(__name__)

app.config.from_pyfile('debug_environment.cfg')

# Get current path
path = os.path.dirname(os.path.abspath(__file__))

app.config["SESSION_FILE_DIR"] = mkdtemp()
app.config["SESSION_PERMANENT"] = True
app.config["SESSION_TYPE"] = "filesystem"
Session(app)
bcrypt = Bcrypt(app)

INPUT_FOLDER = os.path.join(path, 'static/template_file/')
TRASH_FOLDER = os.path.join(path, 'trash')
app.config['INPUT_FOLDER'] = INPUT_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 2147483648  # 2GB limit
app.config['TRASH_FOLDER'] = TRASH_FOLDER


# Allowed extension you can set your own
ALLOWED_EXTENSIONS = set(['tgz'])
IMAGES_EXTENSIONS = set(['png', 'jpg'])

login = False

if login:
    UPLOAD_FOLDER = os.path.join(os.path.join(path, f"uploads{str(session['username'])}"))
    OUTPUT_FOLDER = os.path.join(os.path.join(path, "output"))
    app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
    app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER


# function to add
@app.after_request
def after_request(response):
    response.headers["Cache-Control"] = "no-cache"
    response.headers["Expires"] = 0
    response.headers["Pragma"] = "no-cache"
    response.headers["Referrer-Policy"] = "no-referrer"
    return response


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def allowed_image_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in IMAGES_EXTENSIONS


def login_required(f):
    """
    Decorate routes to require login.
    https://flask.palletsprojects.com/en/1.1.x/patterns/viewdecorators/
    """
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if session.get("username") is None:
            return redirect("/login")
        return f(*args, **kwargs)
    return decorated_function


def admin_required(f):
    """
    Decorate routes to require login.
    https://flask.palletsprojects.com/en/1.1.x/patterns/viewdecorators/
    """
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if session.get("usertype") == 1:
            return redirect("/select_template")
        return f(*args, **kwargs)
    return decorated_function


@app.route('/login', methods=['GET', 'POST'])
def login_info():
    global login
    try:
        if session.get("report_name") is not None:
            if os.path.exists(os.path.join(path, "output/" + session["report_name"] + ".docx")):
                os.remove(os.path.join(path, "output/" + session["report_name"] + ".docx"))
        if session.get("username") is not None:
            if os.path.isdir(os.path.join(path, f"uploads{str(session['username'])}")):
                shutil.rmtree(os.path.join(path, f"uploads{str(session['username'])}"))
    except OSError as e:
        print("Error: %s - %s." % (e.filename, e.strerror))
    if request.method == 'POST':
        if request.form['user'] == 'guest':
            session["username"] = secrets.token_urlsafe(8)
            session["usertype"] = 1
            login = True
            return redirect(url_for("home"))
        elif request.form['user'] == 'admin':
            session["username"] = secrets.token_urlsafe(8)
            session["usertype"] = 2
            login = True
            return render_template("login.html", username=session["username"])
    return render_template("entrypage.html")


@app.route('/password', methods=['GET', 'POST'])
def check_password():
    session["password"] = None
    if request.method == 'POST':
        user_password = bcrypt.generate_password_hash('vrx1v1vd')
        session["password"] = user_password
        if bcrypt.check_password_hash(session["password"], request.form['password']):
            return redirect(url_for("home"))
        else:
            error = "Password Incorrect"
            return render_template("login.html", error=error)
    return render_template("login.html")


@app.route("/logout")
@login_required
def logout():
    global login
    try:
        if session.get("report_name") is not None:
            if os.path.exists(os.path.join(path, "output/" + session["report_name"] + ".docx")):
                os.remove(os.path.join(path, "output/" + session["report_name"] + ".docx"))
        if session.get("username") is not None:
            if os.path.isdir(os.path.join(path, f"uploads{str(session['username'])}")):
                shutil.rmtree(os.path.join(path, f"uploads{str(session['username'])}"))
    except OSError as e:
        print("Error: %s - %s." % (e.filename, e.strerror))
    session.clear()
    login = False
    return redirect(url_for("login_info"))


@app.route('/', methods=['GET', 'POST'])
@login_required
@admin_required
def home():
    if request.method == 'POST':
        global ck_title
        ck_title = request.form.get('title')
        return url_for('create_template', ck_title='ck_title')
    else:
        try:
            if session.get("report_name") is not None:
                if os.path.exists(os.path.join(path, "output/" + session["report_name"] + ".docx")):
                    os.remove(os.path.join(path, "output/" + session["report_name"] + ".docx"))
            if session.get("username") is not None:
                if os.path.isdir(os.path.join(path, f"uploads{str(session['username'])}")):
                    shutil.rmtree(os.path.join(path, f"uploads{str(session['username'])}"))
        except OSError as e:
            print("Error: %s - %s." % (e.filename, e.strerror))
        if not session.get("username"):
            return redirect("/login")
        # print("Here is the username " + session["username"])
        os.chdir(INPUT_FOLDER)
        files = glob.glob("*.docx")
        return render_template('homepageCumTemplates.html', files=files, username=session["username"])


@app.route('/streaming', methods=['GET', 'POST'])
@login_required
def streaming():
    return render_template('troubleshooting.html')


@app.route('/upload_docx', methods=['GET', 'POST'])
@login_required
@admin_required
def upload_docx():
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part')
            return render_template('upload_docx.html')
        file = request.files['file']
        if file.filename == '':
            flash('No file selected for uploading')
            return render_template('upload_docx.html')
        docxfile = secure_filename(file.filename)
        if not os.path.isdir(INPUT_FOLDER):
            os.mkdir(INPUT_FOLDER)
        file.save(os.path.join(app.config['INPUT_FOLDER'], docxfile))
        flash('File(s) successfully uploaded')
        return redirect('/')
    else:
        return render_template('upload_docx.html')


@app.route('/create_template/<ck_title>', methods=['GET', 'POST'])
@login_required
@admin_required
def create_template(ck_title):
    if request.method == 'POST':
        global ck_data
        ck_data = request.form.get('editor')
        if ck_data is not None:
            new_parser = HtmlToDocx()
            docx = new_parser.parse_html_string(ck_data)
            docx.save(os.path.join(app.config['INPUT_FOLDER'], ck_title+'.docx'))
            return redirect('/')
    return render_template('create_template.html')


@app.route('/trash')
@login_required
@admin_required
def trash():
    if not os.path.isdir(TRASH_FOLDER):
        os.mkdir(TRASH_FOLDER)
    os.chdir(TRASH_FOLDER)
    files = glob.glob("*.docx")
    return render_template('trash.html', files=files)


@app.route('/edit/<filename>', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_file(filename):
    file_path = os.path.join(INPUT_FOLDER, filename)
    if request.method == 'POST':
        ck_data = request.form.get('editor')
        if ck_data is not None:
            new_parser = HtmlToDocx()
            docx = new_parser.parse_html_string(ck_data)
            docx.save(os.path.join(app.config['INPUT_FOLDER'], file_path))
            return redirect('/')
    else:
        with open(file_path, "rb") as file1:
            result = mammoth.convert_to_html(file1)
        return render_template('create_template.html', result=result.value)


@app.route('/delete/<filename>')
@login_required
@admin_required
def delete_file(filename):
    file_path = os.path.join(INPUT_FOLDER, filename)
    trash_file = os.path.join(TRASH_FOLDER, filename)
    destination = os.path.join(TRASH_FOLDER)
    if os.path.exists(file_path):
        shutil.copy(file_path, destination)
        os.remove(file_path)
        flash("File is successfully deleted")
        return redirect('/')
    elif os.path.exists(trash_file):
        os.remove(trash_file)
        return redirect('/trash')
    else:
        flash("File is not successfully deleted because it is not present or it is deleted")
        return redirect('/')


@app.route('/restore/<filename>')
@login_required
@admin_required
def restore(filename):
    file_path = os.path.join(TRASH_FOLDER, filename)
    destination = os.path.join(INPUT_FOLDER)
    if os.path.exists(file_path):
        shutil.copy(file_path, destination)
        os.remove(file_path)
    else:
        flash("file does not exist")
    return render_template('trash.html')


@app.route('/stream')
@login_required
def stream():
    filepath = os.path.join(path, 'SwiftDoc.log')
    response = Response(open(filepath, "r"), mimetype="text/plain")
    return response


@app.route('/logs')
@login_required
def logs():
    filepath = os.path.join(path, 'SwiftDoc.log')
    return send_file(filepath, as_attachment=True)


@app.route('/select_template', methods=["GET", "POST"])
@login_required
def template():
    if request.method == 'POST':
        global template1, product_type
        session['template1'] = request.form.get("template_select")
        session['product_type'] = request.form.get("product_select")
        logger.info(str(session['template1']))
        return redirect(url_for('upload_file'))
    else:
        try:
            if session.get("report_name") is not None:
                if os.path.exists(os.path.join(path, "output/" + session["report_name"] + ".docx")):
                    os.remove(os.path.join(path, "output/" + session["report_name"] + ".docx"))
            if session.get("username") is not None:
                if os.path.isdir(os.path.join(path, f"uploads{str(session['username'])}")):
                    shutil.rmtree(os.path.join(path, f"uploads{str(session['username'])}"))
        except OSError as e:
            print("Error: %s - %s." % (e.filename, e.strerror))
        os.chdir(INPUT_FOLDER)
        # print("Here is the username " + session["username"])
        files = glob.glob("*.docx")
        return render_template('select_template.html', files=files)


@app.route('/tsf_info', methods=['GET', 'POST'])
@login_required
def upload_file():
    if request.method == 'POST':
        file = request.files['file']
        if file.filename == '':
            flash('No file selected for uploading')
            return redirect(request.url)
        tsffilename = secure_filename(os.path.splitext(file.filename)[0]+session["username"]+'.tgz')
        save_path = os.path.join(os.path.join(path, f"uploads{str(session['username'])}"), tsffilename)
        current_chunk = int(request.form['dzchunkindex'])
        if os.path.exists(save_path) and current_chunk == 0:
            return make_response(('File already exists', 400))
        if not os.path.isdir(os.path.join(path, f"uploads{str(session['username'])}")):
            os.mkdir(os.path.join(path, f"uploads{str(session['username'])}"))
        try:
            with open(save_path, 'ab') as f:
                f.seek(int(request.form['dzchunkbyteoffset']))
                f.write(file.stream.read())
        except OSError:
            # log.exception will include the traceback so we can see what's wrong
            logger.info('Could not write to file')
            return make_response(("Not sure why,"
                                  " but we couldn't write the file to disk", 500))
        total_chunks = int(request.form['dztotalchunkcount'])
        if current_chunk == total_chunks:
            logger.info("file successfully uploaded")
            return render_template("customer_info.html")
        else:
            return str(current_chunk)
    else:
        return render_template('upload_tsf_file.html')


@app.route('/cust_info', methods=['GET', 'POST'])
@login_required
def cust_info():
    if request.method == 'POST':
        session['customer_name'] = request.form['customername']
        if session['customer_name'] == "":
            flash("Please enter customer name")
            return redirect(url_for('cust_info'))
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)
        logo = request.files['file']
        img = Image.open(logo)
        json_data = json.dumps(np.array(img).tolist())
        session['new_image'] = Image.fromarray(np.array(json.loads(json_data), dtype='uint8'))
        session['imagename'] = secure_filename(logo.filename)
        session['new_image'].save(os.path.join(os.path.join(path, f"uploads{str(session['username'])}"), session['imagename']))
        # message1 = f"Your customer logo({logo.filename}) is successfully uploaded"
        # message2 = f"Your customer name is: {session['customer_name']}"
        return redirect('/report')
    else:
        return render_template('customer_info.html')


@app.route('/report', methods=['GET', 'POST'])
@login_required
def report():
    if request.method == 'POST':
        session['report_name'] = request.form['reportname']
        if session['report_name'] == "":
            flash("Please enter report name")
            return redirect(request.url)
        elif len(session['report_name']) > 35:
            flash("The length of report name should not be greater than 35 characters.")
            return redirect(request.url)
        elif session['report_name'] != "" and len(session['report_name']) <= 35:
            if not os.path.isdir(os.path.join(path, "output")):
                os.mkdir(os.path.join(path, "output"))
            psmain()  # Starting the Report Generator in the Backend
            docx_file = session['report_name'] + ".docx"
            doc = Document(os.path.join(os.path.join(path, "output"), docx_file))
            header = doc.sections[0].first_page_header
            image1 = os.path.join(app.config['INPUT_FOLDER'], 'Picture1.png')
            image2 = os.path.join(os.path.join(path, f"uploads{str(session['username'])}"), session['imagename'])
            header_tp = header.add_paragraph()
            header_run = header_tp.add_run()
            header_run.add_picture(image2, width=Inches(2.051181), height=Inches(0.3740157))
            header_run.add_text('\t\t\t\t')
            header_run.add_picture(image1, width=Inches(2.051181), height=Inches(0.3740157))
            reportfile = os.path.join(os.path.join(path, "output"), session['report_name']) + '.docx'
            doc.save(reportfile)
            report_status = True
            try:
                if report_status and os.path.isdir(os.path.join(path, f"uploads{str(session['username'])}")):
                    shutil.rmtree(os.path.join(path, f"uploads{str(session['username'])}"))
                    print("Successfully removed all files")
            except OSError as e:
                print("Error: %s - %s." % (e.filename, e.strerror))
            return str("success")
        else:
            flash("Please enter valid report name")
            return redirect(request.url)
    else:
        return render_template('generate_report.html')


@app.route('/download')
@login_required
def download():
    # filepath = os.path.join(app.config['OUTPUT_FOLDER'], report_name) + '.docx'
    if not os.path.exists(os.path.join(os.path.join(path, "output"), session['report_name']) + '.docx'):
        logger.info("File not found")
        flash("File not generated properly")
    session['filepath'] = os.path.join(os.path.join(path, "output"), session['report_name']) + '.docx'
    return send_file(session['filepath'], as_attachment=True)


@app.route('/preview')
@login_required
def preview():
    with open(os.path.join(os.path.join(path, "output"), session['report_name']) + '.docx', "rb") as file:
        result = mammoth.convert_to_html(file)
        html_file = result.value
    return html_file


def psmain():
    process_tsf = False
    psusername = str(session["username"])
    my_helper = PanHelpers()
    logger.info('Starting "SwiftDoc" As Built Generator')
    my_helper.print_output('Starting "SwiftDoc" As Built Generator')
    my_helper.debug_info()
    my_helper.print_output("Setting paths")
    my_helper.date_string = datetime.datetime.now().strftime(
        "%y%m%d-%H%M%S"
    )  # used for naming the temp directory
    my_helper.set_separator()  # slash or backslash? This is your legacy Bill Gates, hope you're happy.
    my_helper.set_path(psusername)
    context = {'Customer': session['customer_name'], 'Month': datetime.datetime.now().strftime('%B'),

               'Year': datetime.datetime.now().strftime('%Y')}
    pan_hosts = []
    # Declare an empty list to hold the devices we find TSF and XML creds for.
    """Process the provided TSF files."""
    my_helper.print_output("Processing your TSF files, just a moment please...")
    my_tsf_handler = PanTSFHandler(
        my_helper.input_path, my_helper.tmp_path, my_helper.my_separator
    )  # Create a TSF handler object
    customer_files = my_tsf_handler.load_tsf_archive()
    if customer_files is not None:  # see if there are ay tgz files in inputs/ directory
        logger.debug("Set process_tsf to True.")
        process_tsf = True
    else:
        logger.info("No TSF files found in inputs directory.")
    if process_tsf:
        for my_file in customer_files:
            my_helper.set_path(psusername)  # make sure the temp path exists
            logger.info("Processing file: %s", my_file)
            tsf_info = my_tsf_handler.extract_info_from_tsf(my_file)
            my_helper.remove_tmp_dir()  # blow away the temp files for this device

            hostname = tsf_info["hostname"]
            serial = tsf_info["serial"]

            if tsf_info["model"] in ["Panorama", "M-100", "M-500", "M-200", "M-600"]:
                my_panorama = PaloPanorama(hostname, serial)

                my_panorama.templates_prisma = tsf_info["templates_prisma"]
                my_panorama.opmode = tsf_info["opmode"]
                my_panorama.admins = tsf_info["admins"]
                my_panorama.templates = tsf_info["templates"]
                my_panorama.devicegroups = tsf_info["devicegroups"]
                my_panorama.snmpv2_system = tsf_info["snmpv2s"]
                my_panorama.snmpv2_manager = tsf_info["snmpv2m"]
                my_panorama.snmpv3 = tsf_info["snmpv3"]
                my_panorama.hapriority = tsf_info["hapriority"]
                my_panorama.model = tsf_info["model"]
                my_panorama.panos = tsf_info["panos"]
                my_panorama.licenses = tsf_info["licenses"]
                my_panorama.mgmtinfo = tsf_info["mgmtinfo"]
                pan_hosts.append(my_panorama)
            else:
                my_fw = PaloFirewall(hostname, serial)
                pan_hosts.append(my_fw)

                if "hainfo" in tsf_info:
                    logger.debug("Processing HA info.")
                    my_fw.hainfo = tsf_info["hainfo"]
                    my_fw.haotherinfo = tsf_info["haotherinfo"]
                    my_fw.hapeername = tsf_info["hapeername"]

                my_fw.multivsys = tsf_info["multivsys"]
                my_fw.model = tsf_info["model"]
                my_fw.panos = tsf_info["panos"]
                my_fw.licenses = tsf_info["licenses"]
                my_fw.mgmtinfo = tsf_info["mgmtinfo"]
                my_fw.admins = tsf_info["admins"]
                my_fw.vsys = tsf_info["vsys"]
                my_fw.interfaces = tsf_info["interfaces"]
                my_fw.logintsettings = tsf_info["logintsettings"]
                my_fw.antivirus = tsf_info["antivirus"]
                my_fw.antispy = tsf_info["antispy"]
                my_fw.vulnerability = tsf_info["vulnerability"]
                my_fw.url = tsf_info["url"]
                my_fw.wildfire = tsf_info["wildfire"]
                my_fw.fileblocking = tsf_info["fileblocking"]
                my_fw.update_sched = tsf_info["update_sched"]
                my_fw.data_filtering = tsf_info["data_filtering"]
                my_fw.data_objects = tsf_info["data_objects"]
                my_fw.secprof = tsf_info["secprof"]
                my_fw.dos = tsf_info["dos"]
                my_fw.ike_crypto = tsf_info["ike_crypto"]
                my_fw.ipsec_crypto = tsf_info["ipsec_crypto"]
                my_fw.ike_gw = tsf_info["ike_gw"]
                my_fw.ipsec_vpn = tsf_info["ipsec_vpn"]
                my_fw.zones = tsf_info["zones"]
                my_fw.logforwarding = tsf_info["logforwarding"]
                my_fw.snmpv2_system = tsf_info["snmpv2s"]
                my_fw.snmpv2_manager = tsf_info["snmpv2m"]
                my_fw.snmpv3 = tsf_info["snmpv3"]
                my_fw.gp_portals = tsf_info["gp_portals"]
                my_fw.gp_gateways = tsf_info["gp_gateways"]
                my_fw.hip = tsf_info["hip"]
                my_fw.groups = tsf_info["groups"]
                my_fw.userid = tsf_info["userid"]

                hapeername = my_fw.hapeername
                if hapeername:
                    peerhainfo = {}
                    pan_hosts[hapeername] = PaloFirewall(
                        hapeername, my_fw.hainfo["peerserial"]
                    )
                    pan_hosts[hapeername].model = tsf_info["model"]
                    pan_hosts[hapeername].panos = tsf_info["panos"]
                    pan_hosts[hapeername].licenses = my_fw.licenses
                    peerhainfo["ha1a"] = pan_hosts[-2].hainfo["peerha1a"]
                    peerhainfo["ha1b"] = pan_hosts[-2].hainfo["peerha1b"]
                    peerhainfo["ha2a"] = pan_hosts[-2].hainfo["peerha2a"]
                    peerhainfo["ha2b"] = pan_hosts[-2].hainfo["peerha2b"]
                    peerhainfo["hamode"] = pan_hosts[-2].hainfo["hamode"]
                    peermgmtinfo = {
                        "ip": pan_hosts[-2].hainfo["peermgmtip"],
                        "ipv6": pan_hosts[-2].hainfo["peermgmtipv6"],
                        "mask": pan_hosts[-2].mgmtinfo["mask"],
                        "gw": pan_hosts[-2].mgmtinfo["gw"],
                        "speed": pan_hosts[-2].mgmtinfo["speed"],
                        "mtu": pan_hosts[-2].mgmtinfo["mtu"],
                        "services": pan_hosts[-2].mgmtinfo["services"],
                        "permitted_ips": pan_hosts[-2].mgmtinfo["permitted_ips"],
                    }
                    my_fw.hainfo["heartbeat"] = tsf_info["hainfo"]["heartbeat"]
                    my_fw.mgmtinfo = peermgmtinfo
            logger.info("Finished Processing TS file: %s", my_file)
    else:
        logger.info("No valid TSF files found.")

    if not process_tsf:
        logger.info("Nothing to do!")
        my_helper.print_output("Nothing to do!")
    else:
        logger.info("Generating SwiftDoc.")
        my_helper.print_output("Generating as-built.")

        if session['product_type'] == "Firewall":
            my_firewall_report = Firewall(
                my_helper.date_string,
                my_helper.input_path,
                my_helper.output_path,
                session['report_name'],
                session['template1'],
                context
            )
            my_firewall_report.pan_hosts = pan_hosts
            my_firewall_report.populate_lists()
            my_firewall_report.gen_doc()

        elif session['product_type'] == "Panorama":
            my_panorama_report = Panorama(
                my_helper.date_string,
                my_helper.input_path,
                my_helper.output_path,
                session['report_name'],
                session['template1'],
                context
            )
            my_panorama_report.pan_hosts = pan_hosts
            my_panorama_report.populate_lists()
            my_panorama_report.gen_doc()

        elif session['product_type'] == "All":
            my_swift_doc = SwiftDoc(
                my_helper.date_string,
                my_helper.input_path,
                my_helper.output_path,
                session['report_name'],
                session['template1'],
                context
            )
            my_swift_doc.pan_hosts = pan_hosts
            my_swift_doc.populate_lists()
            my_swift_doc.gen_doc()


if __name__ == "__main__":  # on running python app.py
    # port = int(os.environ.get('PORT', 8080))
    app.run(threaded=True, port=int(os.environ.get('PORT', 8080)))  # run the flask app
